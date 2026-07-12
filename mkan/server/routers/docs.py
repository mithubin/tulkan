import io
import json
import os
import re
import sqlite3
import tempfile
from copy import deepcopy
from datetime import date

from fastapi import APIRouter, Body, Depends, HTTPException
from fastapi.responses import JSONResponse

from auth import current_user_id, require_board_access
from db import get_conn, now, uid
from routers.events import broadcast

router = APIRouter()

UPLOAD_PATH = os.environ.get('UPLOAD_PATH', './data/uploads')

# ── Board-DB helper (parallel zu boarddb.py) ─────────────────────────────────

def _board_db_path(board_id: str) -> str:
    db_path = os.environ.get('DB_PATH', './data/db/kanban.sqlite')
    board_dir = os.path.join(os.path.dirname(os.path.abspath(db_path)), 'boards')
    return os.path.join(board_dir, f'board-{board_id}.db')


def _load_rows(board_id: str, src_card: sqlite3.Row) -> tuple[list[dict], list[str]]:
    """Gibt (rows, columns) aus table_card oder query_card zurück."""
    path = _board_db_path(board_id)
    if not os.path.exists(path):
        raise HTTPException(400, 'Keine Board-Datenbank gefunden')
    bconn = sqlite3.connect(path)
    bconn.row_factory = sqlite3.Row
    try:
        if src_card['card_type'] == 'table_card':
            tname = re.sub(r'[^\w\s\-]', '', src_card['title']).strip()
            cur = bconn.execute(f'SELECT * FROM "{tname}"')
        else:
            sql = (src_card['notes'] or '').strip().rstrip(';')
            if not sql:
                raise HTTPException(400, 'Query-Karte hat keinen SQL-Text')
            cur = bconn.execute(sql)
        columns = [d[0] for d in cur.description or []]
        rows = [dict(zip(columns, row)) for row in cur.fetchall()]
        return rows, columns
    finally:
        bconn.close()


# ── python-docx Platzhalter-Füllung ──────────────────────────────────────────

def _fill_paragraph(para, row: dict) -> None:
    """Setzt Platzhalter auf Paragraphen-Ebene (run-übergreifend sicher)."""
    full = ''.join(r.text for r in para.runs)
    for key, val in row.items():
        full = full.replace('{{' + key + '}}', str(val) if val is not None else '')
    full = full.replace('{{datum}}', date.today().strftime('%d.%m.%Y'))
    if para.runs:
        para.runs[0].text = full
        for r in para.runs[1:]:
            r.text = ''


def _fill_docx(template_path: str, row: dict) -> bytes:
    from docx import Document
    doc = Document(template_path)
    for para in doc.paragraphs:
        _fill_paragraph(para, row)
    for table in doc.tables:
        for tr in table.rows:
            for cell in tr.cells:
                for para in cell.paragraphs:
                    _fill_paragraph(para, row)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_filename(template: str, row: dict, idx: int, ext: str) -> str:
    if not template:
        return f'{idx + 1:03d}_dokument.{ext}'
    name = template
    for key, val in row.items():
        name = name.replace('{{' + key + '}}', re.sub(r'[^\w\-.]', '_', str(val or '')))
    # Endung erzwingen – vorhandene Endung immer entfernen (egal welche)
    base = re.sub(r'\.[a-zA-Z0-9]{1,6}$', '', name)
    return f'{base}.{ext}'


def _safe_filename(name: str) -> str:
    return re.sub(r'[^\w\-.]', '_', name)[:120]


# ── Endpoint: Vorlage anlegen ─────────────────────────────────────────────────

@router.post('/{card_id}/doc-template', status_code=201)
def create_doc_template(card_id: str, user_id: str = Depends(current_user_id)):
    with get_conn() as conn:
        parent = conn.execute('SELECT * FROM cards WHERE id=?', (card_id,)).fetchone()
        if not parent:
            raise HTTPException(404, 'Karte nicht gefunden')
        require_board_access(parent['board_id'], user_id, conn)

        # file_card anlegen
        fc_id = uid()
        ts = now()
        pos_row = conn.execute(
            'SELECT MAX(position) AS m FROM cards WHERE parent_card_id=?', (card_id,)
        ).fetchone()
        pos = (pos_row['m'] or 0) + 1
        conn.execute(
            """INSERT INTO cards
               (id, board_id, col_id, lane_id, position, title, notes,
                color, bg_color, cover_path, points, points_max,
                card_type, parent_card_id, card_mode, time_spent, created_at, updated_at, created_by)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (fc_id, parent['board_id'], parent['col_id'], parent['lane_id'], pos,
             'vorlage.docx', '', None, None, None, 0, None,
             'file_card', card_id, 'org', 0, ts, ts, user_id),
        )

        # Minimale DOCX-Datei anlegen
        from docx import Document
        doc = Document()
        doc.add_paragraph('Hallo {{name}},')
        doc.add_paragraph('')
        doc.add_paragraph('hier kommt dein Text…')
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        card_dir = os.path.join(UPLOAD_PATH, fc_id)
        os.makedirs(card_dir, exist_ok=True)
        att_id = uid()
        filename = 'vorlage.docx'
        path = os.path.join(card_dir, f'{att_id}_{filename}')
        with open(path, 'wb') as f:
            f.write(docx_bytes)

        mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        conn.execute(
            'INSERT INTO attachments (id, card_id, filename, path, size, mime, created_at) VALUES (?,?,?,?,?,?,?)',
            (att_id, fc_id, filename, path, len(docx_bytes), mime, ts),
        )

    broadcast(parent['board_id'], {
        'type': 'card_created', 'userId': user_id, 'boardId': parent['board_id'],
        'id': fc_id, 'title': filename,
        'colId': parent['col_id'], 'laneId': parent['lane_id'],
        'color': None, 'bgColor': None, 'coverImage': None, 'notes': '',
        'subtasks': [], 'attachments': [{
            'id': att_id, 'name': filename, 'size': len(docx_bytes),
            'type': mime, 'url': f'/attachments/{att_id}/file',
        }],
        'labelIds': [], 'points': 0, 'pointsMax': None, 'createdAt': ts,
        'cardType': 'file_card', 'parentCardId': card_id,
    })

    return {
        'cardId': fc_id,
        'card': {
            'id': fc_id, 'title': filename, 'cardType': 'file_card',
            'parentCardId': card_id, 'cardMode': 'org',
            'colId': parent['col_id'], 'laneId': parent['lane_id'],
            'position': pos, 'createdAt': ts,
            'attachments': [{'id': att_id, 'name': filename, 'size': len(docx_bytes),
                              'type': mime, 'url': f'/attachments/{att_id}/file'}],
        },
    }


# ── Endpoint: Seriendokumente erstellen ───────────────────────────────────────

@router.post('/{card_id}/series-docs')
def series_docs(card_id: str, body: dict = Body(...), user_id: str = Depends(current_user_id)):
    output_format = body.get('outputFormat', 'docx')

    with get_conn() as conn:
        card = conn.execute('SELECT * FROM cards WHERE id=?', (card_id,)).fetchone()
        if not card:
            raise HTTPException(404, 'Karte nicht gefunden')
        require_board_access(card['board_id'], user_id, conn)

        cs = json.loads(card['card_settings'] or '{}')
        tpl_card_id = cs.get('templateCardId')
        data_src_id = cs.get('dataSrcId')
        filename_tpl = cs.get('filenameTemplate') or ''

        if not tpl_card_id:
            raise HTTPException(400, 'Keine Vorlage verknüpft')
        if not data_src_id:
            raise HTTPException(400, 'Keine Datenquelle gewählt')

        tpl_att = conn.execute(
            'SELECT * FROM attachments WHERE card_id=? ORDER BY created_at LIMIT 1', (tpl_card_id,)
        ).fetchone()
        if not tpl_att:
            raise HTTPException(400, 'Vorlage hat keinen Anhang')

        src_card = conn.execute('SELECT * FROM cards WHERE id=?', (data_src_id,)).fetchone()
        if not src_card:
            raise HTTPException(400, 'Datenquelle nicht gefunden')

        # Daten laden
        rows, columns = _load_rows(card['board_id'], src_card)
        if not rows:
            raise HTTPException(400, 'Datenquelle ist leer')

        # Vorlagen-Titel für Ausgabe-Karte
        tpl_card = conn.execute('SELECT title FROM cards WHERE id=?', (tpl_card_id,)).fetchone()
        tpl_name = re.sub(r'\.(docx|pdf)$', '', tpl_card['title'] if tpl_card else 'Vorlage', flags=re.IGNORECASE)

        # Ausgabe-Karte: Name mit Nummerierung bei Wiederholung
        base_title = f'Ausgabe - {tpl_name}'
        existing = conn.execute(
            "SELECT title FROM cards WHERE board_id=? AND col_id=? AND lane_id=? AND title LIKE ?",
            (card['board_id'], card['col_id'], card['lane_id'], base_title + '%'),
        ).fetchall()
        if not existing:
            out_title = base_title
        else:
            n = len(existing) + 1
            out_title = f'{base_title} ({n})'

        # Ausgabe-Karte anlegen
        out_id = uid()
        ts = now()
        pos_row = conn.execute(
            'SELECT MAX(position) AS m FROM cards WHERE col_id=? AND lane_id=? AND parent_card_id IS NULL',
            (card['col_id'], card['lane_id']),
        ).fetchone()
        out_pos = (pos_row['m'] or 0) + 1

        conn.execute(
            """INSERT INTO cards
               (id, board_id, col_id, lane_id, position, title, notes,
                color, bg_color, cover_path, points, points_max,
                card_type, parent_card_id, card_mode, time_spent, created_at, updated_at, created_by,
                card_settings)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (out_id, card['board_id'], card['col_id'], card['lane_id'], out_pos,
             out_title, '', None, None, None, 0, None,
             'card', None, 'org', 0, ts, ts, user_id,
             json.dumps({'docSourceCardId': card_id})),
        )

        # selectedIndices: None = alle Zeilen; sonst Teilmenge mit orig. Index als position
        selected_indices = body.get('selectedIndices')
        if selected_indices is not None:
            indexed_rows = [(i, rows[i]) for i in selected_indices if 0 <= i < len(rows)]
        else:
            indexed_rows = list(enumerate(rows))
        if not indexed_rows:
            raise HTTPException(400, 'Keine Zeilen ausgewählt')

        # Pro Zeile: DOCX füllen, file_card + Anhang anlegen
        created = []
        use_pdf = output_format == 'pdf'

        for orig_idx, row_data in indexed_rows:
            docx_bytes = _fill_docx(tpl_att['path'], row_data)

            if use_pdf:
                # PDF-Konvertierung via OnlyOffice/convert.py (falls verfügbar)
                try:
                    from routers.convert import _convert_bytes_to_pdf
                    final_bytes = _convert_bytes_to_pdf(docx_bytes, 'docx')
                    final_ext = 'pdf'
                    final_mime = 'application/pdf'
                except Exception:
                    # Fallback: DOCX ausgeben
                    final_bytes = docx_bytes
                    final_ext = 'docx'
                    final_mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                final_bytes = docx_bytes
                final_ext = 'docx'
                final_mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

            fname = _make_filename(filename_tpl, row_data, orig_idx, final_ext)

            # file_card: position = orig_idx → Mail-Versand matcht per origIdx
            fc_id = uid()
            conn.execute(
                """INSERT INTO cards
                   (id, board_id, col_id, lane_id, position, title, notes,
                    color, bg_color, cover_path, points, points_max,
                    card_type, parent_card_id, card_mode, time_spent, created_at, updated_at, created_by)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (fc_id, card['board_id'], card['col_id'], card['lane_id'], orig_idx,
                 fname, '', None, None, None, 0, None,
                 'file_card', out_id, 'org', 0, ts, ts, user_id),
            )

            card_dir = os.path.join(UPLOAD_PATH, fc_id)
            os.makedirs(card_dir, exist_ok=True)
            att_id = uid()
            path = os.path.join(card_dir, f'{att_id}_{_safe_filename(fname)}')
            with open(path, 'wb') as f:
                f.write(final_bytes)
            conn.execute(
                'INSERT INTO attachments (id, card_id, filename, path, size, mime, created_at) VALUES (?,?,?,?,?,?,?)',
                (att_id, fc_id, fname, path, len(final_bytes), final_mime, ts),
            )
            created.append({'fcId': fc_id, 'attId': att_id, 'name': fname})

    broadcast(card['board_id'], {'type': 'board_changed', 'boardId': card['board_id']})
    return {'count': len(created), 'cardId': out_id, 'cardTitle': out_title, 'files': created}


# ── Endpoint: Doclink-Seriengenerierung (aus doclink_card heraus) ─────────────

@router.post('/{card_id}/doclink-series')
def doclink_series(card_id: str, body: dict = Body(...), user_id: str = Depends(current_user_id)):
    """Generiert Seriendokumente aus einer doclink_card.
    Erwartet: templateCardId (doc-mode Karte), dataSrcId, outputFormat, filenameTemplate, selectedIndices.
    templateCardId zeigt auf die board-seitige doc-Karte; deren cardSettings.templateCardId gibt die DOCX-Datei-Karte an.
    """
    template_board_card_id = body.get('templateCardId')
    data_src_id = body.get('dataSrcId')
    output_format = body.get('outputFormat', 'docx')
    filename_tpl = body.get('filenameTemplate') or ''
    selected_indices = body.get('selectedIndices')  # None = alle Zeilen

    with get_conn() as conn:
        card = conn.execute('SELECT * FROM cards WHERE id=?', (card_id,)).fetchone()
        if not card:
            raise HTTPException(404, 'Karte nicht gefunden')
        require_board_access(card['board_id'], user_id, conn)

        if not template_board_card_id:
            raise HTTPException(400, 'Keine Vorlage gewählt')
        if not data_src_id:
            raise HTTPException(400, 'Keine Datenquelle gewählt')

        # doc-mode Karte → deren templateCardId → DOCX-Anhang
        tpl_board_card = conn.execute('SELECT * FROM cards WHERE id=?', (template_board_card_id,)).fetchone()
        if not tpl_board_card:
            raise HTTPException(400, 'Vorlagen-Karte nicht gefunden')
        tpl_cs = json.loads(tpl_board_card['card_settings'] or '{}')
        tpl_file_card_id = tpl_cs.get('templateCardId')
        if not tpl_file_card_id:
            raise HTTPException(400, 'Vorlagen-Karte hat noch kein DOCX (Vorlage anlegen)')
        tpl_att = conn.execute(
            'SELECT * FROM attachments WHERE card_id=? ORDER BY created_at LIMIT 1', (tpl_file_card_id,)
        ).fetchone()
        if not tpl_att:
            raise HTTPException(400, 'DOCX-Anhang nicht gefunden')

        src_card = conn.execute('SELECT * FROM cards WHERE id=?', (data_src_id,)).fetchone()
        if not src_card:
            raise HTTPException(400, 'Datenquelle nicht gefunden')

        rows, columns = _load_rows(card['board_id'], src_card)
        if not rows:
            raise HTTPException(400, 'Datenquelle ist leer')

        if selected_indices is not None:
            indexed_rows = [(i, rows[i]) for i in selected_indices if 0 <= i < len(rows)]
        else:
            indexed_rows = list(enumerate(rows))
        if not indexed_rows:
            raise HTTPException(400, 'Keine Zeilen ausgewählt')

        # Ausgabe-Karte (gleiche Spalte/Lane wie doclink_card)
        tpl_name = re.sub(r'\.(docx|pdf)$', '', tpl_board_card['title'] or 'Vorlage', flags=re.IGNORECASE)
        base_title = f'Ausgabe - {tpl_name}'
        existing = conn.execute(
            "SELECT title FROM cards WHERE board_id=? AND col_id=? AND lane_id=? AND title LIKE ?",
            (card['board_id'], card['col_id'], card['lane_id'], base_title + '%'),
        ).fetchall()
        out_title = base_title if not existing else f'{base_title} ({len(existing) + 1})'

        out_id = uid()
        ts = now()
        pos_row = conn.execute(
            'SELECT MAX(position) AS m FROM cards WHERE col_id=? AND lane_id=? AND parent_card_id IS NULL',
            (card['col_id'], card['lane_id']),
        ).fetchone()
        out_pos = (pos_row['m'] or 0) + 1
        conn.execute(
            """INSERT INTO cards
               (id, board_id, col_id, lane_id, position, title, notes,
                color, bg_color, cover_path, points, points_max,
                card_type, parent_card_id, card_mode, time_spent, created_at, updated_at, created_by,
                card_settings)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (out_id, card['board_id'], card['col_id'], card['lane_id'], out_pos,
             out_title, '', None, None, None, 0, None,
             'card', None, 'org', 0, ts, ts, user_id,
             json.dumps({'docSourceCardId': card_id})),
        )

        created = []
        use_pdf = output_format == 'pdf'

        for orig_idx, row_data in indexed_rows:
            docx_bytes = _fill_docx(tpl_att['path'], row_data)
            if use_pdf:
                try:
                    from routers.convert import _convert_bytes_to_pdf
                    final_bytes = _convert_bytes_to_pdf(docx_bytes, 'docx')
                    final_ext = 'pdf'
                    final_mime = 'application/pdf'
                except Exception:
                    final_bytes = docx_bytes
                    final_ext = 'docx'
                    final_mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                final_bytes = docx_bytes
                final_ext = 'docx'
                final_mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

            fname = _make_filename(filename_tpl, row_data, orig_idx, final_ext)
            fc_id = uid()
            conn.execute(
                """INSERT INTO cards
                   (id, board_id, col_id, lane_id, position, title, notes,
                    color, bg_color, cover_path, points, points_max,
                    card_type, parent_card_id, card_mode, time_spent, created_at, updated_at, created_by)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (fc_id, card['board_id'], card['col_id'], card['lane_id'], orig_idx,
                 fname, '', None, None, None, 0, None,
                 'file_card', out_id, 'org', 0, ts, ts, user_id),
            )
            card_dir = os.path.join(UPLOAD_PATH, fc_id)
            os.makedirs(card_dir, exist_ok=True)
            att_id = uid()
            path = os.path.join(card_dir, f'{att_id}_{_safe_filename(fname)}')
            with open(path, 'wb') as f:
                f.write(final_bytes)
            conn.execute(
                'INSERT INTO attachments (id, card_id, filename, path, size, mime, created_at) VALUES (?,?,?,?,?,?,?)',
                (att_id, fc_id, fname, path, len(final_bytes), final_mime, ts),
            )
            created.append({'fcId': fc_id, 'attId': att_id, 'name': fname})

    broadcast(card['board_id'], {'type': 'board_changed', 'boardId': card['board_id']})
    return {'count': len(created), 'cardId': out_id, 'cardTitle': out_title, 'files': created}


# ── Endpoint: Felder als Hilfszeile ins Dokument schreiben ────────────────────

@router.post('/{card_id}/doc-insert-fields')
def doc_insert_fields(card_id: str, body: dict = Body(...), user_id: str = Depends(current_user_id)):
    """Schreibt alle Felder als erste Zeile ins DOCX-Template (vor OO-Öffnen anwenden)."""
    fields = body.get('fields', [])
    if not fields:
        raise HTTPException(400, 'Keine Felder übergeben')

    with get_conn() as conn:
        card = conn.execute('SELECT * FROM cards WHERE id=?', (card_id,)).fetchone()
        if not card:
            raise HTTPException(404)
        require_board_access(card['board_id'], user_id, conn)

        cs = json.loads(card['card_settings'] or '{}')
        tpl_card_id = cs.get('templateCardId')
        if not tpl_card_id:
            raise HTTPException(400, 'Keine Vorlage verknüpft')

        att = conn.execute(
            'SELECT * FROM attachments WHERE card_id=? ORDER BY created_at LIMIT 1', (tpl_card_id,)
        ).fetchone()
        if not att:
            raise HTTPException(400, 'Vorlage hat keinen Anhang')

    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor

    doc = Document(att['path'])

    # Bestehende Hilfszeile entfernen (erkennen an Marker-Text)
    _MARKER = '‹ Felder:'
    for para in list(doc.paragraphs):
        if para.text.startswith(_MARKER):
            p = para._element
            p.getparent().remove(p)
            break

    # Neue Hilfszeile ganz oben einfügen
    field_text = _MARKER + '  ' + '   '.join('{{' + f + '}}' for f in fields) + '  ›'
    new_para = OxmlElement('w:p')
    new_run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    # Farbe grau
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '888888')
    # Schriftgröße 18 (= 9pt in half-points)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')
    rpr.append(color_el)
    rpr.append(sz)
    new_run.append(rpr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = field_text
    new_run.append(t)
    new_para.append(new_run)
    # Vor den ersten Absatz einfügen
    body_el = doc.element.body
    body_el.insert(0, new_para)

    doc.save(att['path'])

    # Dateigröße aktualisieren
    new_size = os.path.getsize(att['path'])
    with get_conn() as conn:
        conn.execute('UPDATE attachments SET size=? WHERE id=?', (new_size, att['id']))

    return {'ok': True, 'fields': fields}


# ── Endpoint: Seriengenerierung aus serienlink_card ───────────────────────────

@router.post('/{card_id}/serienlink-docs')
def serienlink_docs(card_id: str, body: dict = Body(...), user_id: str = Depends(current_user_id)):
    """Seriengenerierung für serienlink_card.
    Erwartet: dokBoardCardId (doc-mode Boardkarte), dataSrcId, outputFormat, filenameTemplate, selectedIndices.
    file_card.position = origIdx → Mail-Matching per Checkbox-Index.
    """
    dok_board_card_id = body.get('dokBoardCardId')
    data_src_id = body.get('dataSrcId')
    output_format = body.get('outputFormat', 'docx')
    filename_tpl = body.get('filenameTemplate') or ''
    selected_indices = body.get('selectedIndices')

    with get_conn() as conn:
        card = conn.execute('SELECT * FROM cards WHERE id=?', (card_id,)).fetchone()
        if not card:
            raise HTTPException(404, 'Karte nicht gefunden')
        require_board_access(card['board_id'], user_id, conn)

        if not dok_board_card_id:
            raise HTTPException(400, 'Keine Dok-Vorlage gewählt')
        if not data_src_id:
            raise HTTPException(400, 'Keine Datenquelle gewählt')

        tpl_board_card = conn.execute('SELECT * FROM cards WHERE id=?', (dok_board_card_id,)).fetchone()
        if not tpl_board_card:
            raise HTTPException(400, 'Dok-Karte nicht gefunden')
        tpl_cs = json.loads(tpl_board_card['card_settings'] or '{}')
        tpl_file_card_id = tpl_cs.get('templateCardId')
        if not tpl_file_card_id:
            raise HTTPException(400, 'Dok-Karte hat noch kein DOCX (Vorlage anlegen)')
        tpl_att = conn.execute(
            'SELECT * FROM attachments WHERE card_id=? ORDER BY created_at LIMIT 1', (tpl_file_card_id,)
        ).fetchone()
        if not tpl_att:
            raise HTTPException(400, 'DOCX-Anhang nicht gefunden')

        src_card = conn.execute('SELECT * FROM cards WHERE id=?', (data_src_id,)).fetchone()
        if not src_card:
            raise HTTPException(400, 'Datenquelle nicht gefunden')

        rows, columns = _load_rows(card['board_id'], src_card)
        if not rows:
            raise HTTPException(400, 'Datenquelle ist leer')

        if selected_indices is not None:
            indexed_rows = [(i, rows[i]) for i in selected_indices if 0 <= i < len(rows)]
        else:
            indexed_rows = list(enumerate(rows))
        if not indexed_rows:
            raise HTTPException(400, 'Keine Zeilen ausgewählt')

        # Ausgabe-Karte landet bei der Dok-Boardkarte (nicht bei der serienlink_card)
        out_col_id = tpl_board_card['col_id']
        out_lane_id = tpl_board_card['lane_id']

        tpl_name = re.sub(r'\.[a-zA-Z0-9]{1,6}$', '', tpl_board_card['title'] or 'Vorlage')
        base_title = f'Ausgabe - {tpl_name}'
        existing = conn.execute(
            "SELECT title FROM cards WHERE board_id=? AND col_id=? AND lane_id=? AND title LIKE ?",
            (card['board_id'], out_col_id, out_lane_id, base_title + '%'),
        ).fetchall()
        out_title = base_title if not existing else f'{base_title} ({len(existing) + 1})'

        out_id = uid()
        ts = now()
        pos_row = conn.execute(
            'SELECT MAX(position) AS m FROM cards WHERE col_id=? AND lane_id=? AND parent_card_id IS NULL',
            (out_col_id, out_lane_id),
        ).fetchone()
        out_pos = (pos_row['m'] or 0) + 1
        conn.execute(
            """INSERT INTO cards
               (id, board_id, col_id, lane_id, position, title, notes,
                color, bg_color, cover_path, points, points_max,
                card_type, parent_card_id, card_mode, time_spent, created_at, updated_at, created_by,
                card_settings)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (out_id, card['board_id'], out_col_id, out_lane_id, out_pos,
             out_title, '', None, None, None, 0, None,
             'card', None, 'org', 0, ts, ts, user_id,
             json.dumps({'docSourceCardId': card_id})),
        )

        created = []
        use_pdf = output_format == 'pdf'

        for orig_idx, row_data in indexed_rows:
            docx_bytes = _fill_docx(tpl_att['path'], row_data)
            if use_pdf:
                try:
                    from routers.convert import _convert_bytes_to_pdf
                    final_bytes = _convert_bytes_to_pdf(docx_bytes, 'docx')
                    final_ext = 'pdf'
                    final_mime = 'application/pdf'
                except Exception:
                    final_bytes = docx_bytes
                    final_ext = 'docx'
                    final_mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                final_bytes = docx_bytes
                final_ext = 'docx'
                final_mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

            fname = _make_filename(filename_tpl, row_data, orig_idx, final_ext)
            fc_id = uid()
            conn.execute(
                """INSERT INTO cards
                   (id, board_id, col_id, lane_id, position, title, notes,
                    color, bg_color, cover_path, points, points_max,
                    card_type, parent_card_id, card_mode, time_spent, created_at, updated_at, created_by)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (fc_id, card['board_id'], out_col_id, out_lane_id, orig_idx,
                 fname, '', None, None, None, 0, None,
                 'file_card', out_id, 'org', 0, ts, ts, user_id),
            )
            card_dir = os.path.join(UPLOAD_PATH, fc_id)
            os.makedirs(card_dir, exist_ok=True)
            att_id = uid()
            path = os.path.join(card_dir, f'{att_id}_{_safe_filename(fname)}')
            with open(path, 'wb') as f:
                f.write(final_bytes)
            conn.execute(
                'INSERT INTO attachments (id, card_id, filename, path, size, mime, created_at) VALUES (?,?,?,?,?,?,?)',
                (att_id, fc_id, fname, path, len(final_bytes), final_mime, ts),
            )
            created.append({'fcId': fc_id, 'attId': att_id, 'name': fname})

    broadcast(card['board_id'], {'type': 'board_changed', 'boardId': card['board_id']})
    return {'count': len(created), 'cardId': out_id, 'cardTitle': out_title, 'files': created}


# ── Endpoint: Ausgabe-Karten zu einer Quelldok-Karte auflisten ────────────────

@router.get('/{card_id}/ausgabe-cards')
def list_ausgabe_cards(card_id: str, user_id: str = Depends(current_user_id)):
    """Gibt alle Ausgabe-Karten zurück, die aus card_id generiert wurden.
    Matching: card_settings.docSourceCardId == card_id (gespeichert beim Generieren).
    Response: [{id, title, createdAt, fileCount}]
    """
    with get_conn() as conn:
        card = conn.execute('SELECT * FROM cards WHERE id=?', (card_id,)).fetchone()
        if not card:
            raise HTTPException(404, 'Karte nicht gefunden')
        require_board_access(card['board_id'], user_id, conn)

        all_cards = conn.execute(
            "SELECT id, title, created_at, card_settings FROM cards WHERE board_id=? AND card_type='card' AND card_settings IS NOT NULL",
            (card['board_id'],),
        ).fetchall()

        result = []
        for c in all_cards:
            try:
                cs = json.loads(c['card_settings'] or '{}')
            except Exception:
                continue
            if cs.get('docSourceCardId') != card_id:
                continue

            # file_cards with their first attachment, ordered by position
            fc_rows = conn.execute(
                """SELECT fc.id AS fc_id, fc.position, a.id AS att_id, a.filename
                   FROM cards fc
                   LEFT JOIN attachments a ON a.card_id=fc.id
                   WHERE fc.parent_card_id=? AND fc.card_type='file_card'
                   ORDER BY fc.position, a.created_at""",
                (c['id'],),
            ).fetchall()

            # de-duplicate: keep first attachment per file_card
            seen = {}
            files = []
            for row in fc_rows:
                if row['fc_id'] not in seen:
                    seen[row['fc_id']] = True
                    if row['att_id']:
                        files.append({'fcId': row['fc_id'], 'attId': row['att_id'],
                                      'name': row['filename'], 'position': row['position']})

            result.append({
                'id': c['id'],
                'title': c['title'],
                'createdAt': c['created_at'],
                'fileCount': len(files),
                'files': files,
            })

        # Neueste zuerst
        result.sort(key=lambda x: x['createdAt'], reverse=True)
        return result
