import hmac
import os
import re
import tarfile
import tempfile
import zipfile
from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Request, UploadFile, File
from fastapi.responses import FileResponse
from pydantic import BaseModel

from db import get_conn, uid, now
from auth import current_user_id, require_board_access
from routers.events import broadcast

router = APIRouter()

UPLOAD_PATH = os.environ.get('UPLOAD_PATH', './data/uploads')
MAX_SIZE = 1024 * 1024 * 1024  # 1 GB


async def _stream_to_file(file: UploadFile, path: str) -> int:
    size = 0
    try:
        with open(path, 'wb') as f:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                size += len(chunk)
                if size > MAX_SIZE:
                    raise HTTPException(413, 'File too large (max 1 GB)')
                f.write(chunk)
    except HTTPException:
        if os.path.exists(path):
            os.unlink(path)
        raise
    return size


@router.post('/cards/{card_id}', status_code=201)
async def upload_attachment(
    card_id: str,
    file: UploadFile = File(...),
    user_id: str = Depends(current_user_id),
):
    _board_id = None
    with get_conn() as conn:
        c = conn.execute('SELECT board_id FROM cards WHERE id=?', (card_id,)).fetchone()
        if not c:
            raise HTTPException(404, 'Card not found')
        require_board_access(c['board_id'], user_id, conn)
        _board_id = c['board_id']

    card_dir = os.path.join(UPLOAD_PATH, card_id)
    os.makedirs(card_dir, exist_ok=True)

    safe_name = _safe_filename(file.filename or 'upload')
    att_id = uid()
    path = os.path.join(card_dir, f'{att_id}_{safe_name}')
    size = await _stream_to_file(file, path)

    with get_conn() as conn:
        conn.execute(
            'INSERT INTO attachments (id, card_id, filename, path, size, mime, created_at) VALUES (?,?,?,?,?,?,?)',
            (att_id, card_id, file.filename or safe_name, path, size, file.content_type, now()),
        )
    broadcast(_board_id, {'type': 'card_updated', 'userId': user_id, 'id': card_id, 'boardId': _board_id})
    return {
        'id': att_id,
        'name': file.filename,
        'size': size,
        'type': file.content_type,
        'url': f'/attachments/{att_id}/file',
    }


@router.get('/{att_id}/file')
def download_attachment(att_id: str, user_id: str = Depends(current_user_id)):
    with get_conn() as conn:
        att = conn.execute(
            'SELECT a.*, c.board_id FROM attachments a JOIN cards c ON a.card_id=c.id WHERE a.id=?',
            (att_id,),
        ).fetchone()
        if not att:
            raise HTTPException(404, 'Attachment not found')
        require_board_access(att['board_id'], user_id, conn, min_role='viewer')

    if not os.path.exists(att['path']):
        raise HTTPException(404, 'File missing on disk')
    return FileResponse(att['path'], filename=att['filename'], media_type=att['mime'] or 'application/octet-stream')


@router.delete('/{att_id}')
def delete_attachment(att_id: str, user_id: str = Depends(current_user_id)):
    _board_id = _card_id = None
    with get_conn() as conn:
        att = conn.execute(
            'SELECT a.*, c.board_id FROM attachments a JOIN cards c ON a.card_id=c.id WHERE a.id=?',
            (att_id,),
        ).fetchone()
        if not att:
            raise HTTPException(404, 'Attachment not found')
        require_board_access(att['board_id'], user_id, conn)
        _board_id, _card_id = att['board_id'], att['card_id']
        conn.execute('DELETE FROM attachments WHERE id=?', (att_id,))

    try:
        os.remove(att['path'])
    except FileNotFoundError:
        pass
    broadcast(_board_id, {'type': 'card_updated', 'userId': user_id, 'id': _card_id, 'boardId': _board_id})
    return {'ok': True}


# ── Cover-Bild ────────────────────────────────────────────────────────────────

@router.post('/cover/{card_id}', status_code=201)
async def upload_cover(
    card_id: str,
    file: UploadFile = File(...),
    user_id: str = Depends(current_user_id),
):
    _board_id = None
    with get_conn() as conn:
        c = conn.execute('SELECT * FROM cards WHERE id=?', (card_id,)).fetchone()
        if not c:
            raise HTTPException(404, 'Card not found')
        require_board_access(c['board_id'], user_id, conn)
        _board_id = c['board_id']

    card_dir = os.path.join(UPLOAD_PATH, card_id)
    os.makedirs(card_dir, exist_ok=True)
    path = os.path.join(card_dir, 'cover')
    await _stream_to_file(file, path)

    with get_conn() as conn:
        conn.execute('UPDATE cards SET cover_path=?, updated_at=? WHERE id=?', (path, now(), card_id))
    broadcast(_board_id, {'type': 'card_updated', 'userId': user_id, 'id': card_id, 'boardId': _board_id})
    return {'url': f'/attachments/cover/{card_id}'}


@router.get('/cover/{card_id}')
def get_cover(card_id: str, user_id: str = Depends(current_user_id)):
    with get_conn() as conn:
        c = conn.execute('SELECT cover_path, board_id FROM cards WHERE id=?', (card_id,)).fetchone()
        if not c or not c['cover_path']:
            raise HTTPException(404, 'No cover image')
        require_board_access(c['board_id'], user_id, conn, min_role='viewer')

    if not os.path.exists(c['cover_path']):
        raise HTTPException(404, 'File missing on disk')
    import mimetypes
    mime, _ = mimetypes.guess_type(c['cover_path'])
    return FileResponse(c['cover_path'], media_type=mime or 'image/jpeg')


@router.delete('/cover/{card_id}')
def delete_cover(card_id: str, user_id: str = Depends(current_user_id)):
    _board_id = None
    with get_conn() as conn:
        c = conn.execute('SELECT * FROM cards WHERE id=?', (card_id,)).fetchone()
        if not c:
            raise HTTPException(404, 'Card not found')
        require_board_access(c['board_id'], user_id, conn)
        _board_id = c['board_id']
        path = c['cover_path']
        conn.execute('UPDATE cards SET cover_path=NULL, updated_at=? WHERE id=?', (now(), card_id))

    if path:
        try:
            os.remove(path)
        except FileNotFoundError:
            pass
    broadcast(_board_id, {'type': 'card_updated', 'userId': user_id, 'id': card_id, 'boardId': _board_id})
    return {'ok': True}


def _safe_filename(name: str) -> str:
    return re.sub(r'[^\w.\-]', '_', name)[:100]


# ── Card-Attachments ZIP ──────────────────────────────────────────────────────

@router.get('/cards/{card_id}/attachments.zip')
def download_card_attachments_zip(
    card_id: str,
    background_tasks: BackgroundTasks,
    user_id: str = Depends(current_user_id),
):
    with get_conn() as conn:
        card = conn.execute('SELECT id, board_id FROM cards WHERE id=?', (card_id,)).fetchone()
        if not card:
            raise HTTPException(404, 'Card not found')
        require_board_access(card['board_id'], user_id, conn, min_role='viewer')
        atts = [
            {'filename': r['filename'], 'path': r['path']}
            for r in conn.execute(
                'SELECT filename, path FROM attachments WHERE card_id=?', (card_id,)
            ).fetchall()
        ]
        fc_ids = [
            r['id'] for r in conn.execute(
                "SELECT id FROM cards WHERE parent_card_id=? AND card_type='file_card'", (card_id,)
            ).fetchall()
        ]
        for fcid in fc_ids:
            for r in conn.execute('SELECT filename, path FROM attachments WHERE card_id=?', (fcid,)).fetchall():
                atts.append({'filename': r['filename'], 'path': r['path']})

    if not atts:
        raise HTTPException(404, 'Keine Anhänge vorhanden')

    tmp = tempfile.NamedTemporaryFile(suffix='.zip', delete=False)
    try:
        with zipfile.ZipFile(tmp.name, 'w', zipfile.ZIP_DEFLATED) as zf:
            seen: dict[str, int] = {}
            for att in atts:
                if not os.path.exists(att['path']):
                    continue
                name = att['filename']
                if name in seen:
                    seen[name] += 1
                    base, ext = os.path.splitext(att['filename'])
                    name = f'{base}_{seen[name]}{ext}'
                else:
                    seen[name] = 0
                zf.write(att['path'], name)
        tmp.close()
    except Exception:
        tmp.close()
        os.unlink(tmp.name)
        raise HTTPException(500, 'ZIP-Erstellung fehlgeschlagen')

    background_tasks.add_task(os.unlink, tmp.name)
    return FileResponse(tmp.name, media_type='application/zip', filename='attachments.zip')


# ── Attachment verschieben ────────────────────────────────────────────────────

class MoveAttachmentIn(BaseModel):
    card_id: str


class PositionIn(BaseModel):
    position: int


@router.patch('/{att_id}/position')
def set_attachment_position(att_id: str, body: PositionIn, user_id: str = Depends(current_user_id)):
    with get_conn() as conn:
        att = conn.execute(
            'SELECT a.*, c.board_id FROM attachments a JOIN cards c ON a.card_id=c.id WHERE a.id=?',
            (att_id,),
        ).fetchone()
        if not att:
            raise HTTPException(404, 'Attachment not found')
        require_board_access(att['board_id'], user_id, conn)
        conn.execute('UPDATE attachments SET position=? WHERE id=?', (body.position, att_id))
    return {'ok': True}


@router.put('/{att_id}/content')
async def update_attachment_content(
    att_id: str,
    file: UploadFile = File(...),
    user_id: str = Depends(current_user_id),
):
    with get_conn() as conn:
        att = conn.execute(
            'SELECT a.id, a.path, a.card_id, c.board_id FROM attachments a JOIN cards c ON a.card_id=c.id WHERE a.id=?',
            (att_id,),
        ).fetchone()
        if not att:
            raise HTTPException(404, 'Attachment not found')
        require_board_access(att['board_id'], user_id, conn)

    size = await _stream_to_file(file, att['path'])

    with get_conn() as conn:
        conn.execute('UPDATE attachments SET size=? WHERE id=?', (size, att_id))
        conn.execute('UPDATE cards SET updated_at=? WHERE id=?', (now(), att['card_id']))

    broadcast(att['board_id'], {'type': 'card_updated', 'userId': user_id, 'id': att['card_id'], 'boardId': att['board_id']})
    return {'ok': True, 'size': size}


@router.get('/{att_id}/preview-text')
def preview_attachment_text(att_id: str, user_id: str = Depends(current_user_id)):
    with get_conn() as conn:
        att = conn.execute(
            'SELECT a.id, a.filename, a.path, a.card_id, c.board_id FROM attachments a JOIN cards c ON a.card_id=c.id WHERE a.id=?',
            (att_id,),
        ).fetchone()
        if not att:
            raise HTTPException(404, 'Attachment not found')
        require_board_access(att['board_id'], user_id, conn)

    path = att['path']
    name = att['filename'] or ''
    ext = name.rsplit('.', 1)[-1].lower() if '.' in name else ''
    MAX_CHARS = 3000

    try:
        if ext == 'docx':
            from docx import Document
            doc = Document(path)
            text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())

        elif ext == 'odt':
            import xml.etree.ElementTree as ET
            NS_T = 'urn:oasis:names:tc:opendocument:xmlns:text:1.0'
            with zipfile.ZipFile(path) as z:
                with z.open('content.xml') as f:
                    tree = ET.parse(f)
            parts = []
            for p in tree.iter(f'{{{NS_T}}}p'):
                t = ''.join(p.itertext())
                if t.strip():
                    parts.append(t)
            text = '\n'.join(parts)

        elif ext == 'xlsx':
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            ws = wb.active
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else '' for c in row]
                if any(c.strip() for c in cells):
                    rows.append('\t'.join(cells))
                if sum(len(r) for r in rows) > MAX_CHARS:
                    break
            wb.close()
            text = '\n'.join(rows)

        elif ext == 'ods':
            import xml.etree.ElementTree as ET
            NS_TAB = 'urn:oasis:names:tc:opendocument:xmlns:table:1.0'
            with zipfile.ZipFile(path) as z:
                with z.open('content.xml') as f:
                    tree = ET.parse(f)
            rows = []
            for row in tree.iter(f'{{{NS_TAB}}}table-row'):
                cells = [''.join(cell.itertext()).strip()
                         for cell in row.iter(f'{{{NS_TAB}}}table-cell')]
                line = '\t'.join(cells).rstrip('\t')
                if line.strip():
                    rows.append(line)
                if sum(len(r) for r in rows) > MAX_CHARS:
                    break
            text = '\n'.join(rows)

        elif ext == 'pptx':
            import xml.etree.ElementTree as ET
            NS_A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
            parts = []
            with zipfile.ZipFile(path) as z:
                slides = sorted(n for n in z.namelist()
                                if n.startswith('ppt/slides/slide') and n.endswith('.xml'))
                for sname in slides:
                    with z.open(sname) as f:
                        tree = ET.parse(f)
                    texts = [t.text for t in tree.iter(f'{{{NS_A}}}t')
                             if t.text and t.text.strip()]
                    if texts:
                        parts.append(' '.join(texts))
                    if sum(len(p) for p in parts) > MAX_CHARS:
                        break
            text = '\n'.join(parts)

        elif ext == 'odp':
            import xml.etree.ElementTree as ET
            NS_T = 'urn:oasis:names:tc:opendocument:xmlns:text:1.0'
            with zipfile.ZipFile(path) as z:
                with z.open('content.xml') as f:
                    tree = ET.parse(f)
            parts = []
            for p in tree.iter(f'{{{NS_T}}}p'):
                t = ''.join(p.itertext())
                if t.strip():
                    parts.append(t)
            text = '\n'.join(parts)

        elif ext == 'xls':
            import xlrd
            wb = xlrd.open_workbook(path)
            ws = wb.sheet_by_index(0)
            rows = []
            for i in range(ws.nrows):
                cells = [str(ws.cell_value(i, j)) for j in range(ws.ncols)]
                line = '\t'.join(c for c in cells if c)
                if line.strip():
                    rows.append(line)
                if sum(len(r) for r in rows) > MAX_CHARS:
                    break
            text = '\n'.join(rows)

        elif ext == 'doc':
            import subprocess
            result = subprocess.run(['antiword', path], capture_output=True, timeout=15)
            if result.returncode != 0:
                raise HTTPException(400, 'antiword nicht verfügbar oder Lesefehler')
            text = result.stdout.decode('utf-8', errors='replace')

        else:
            raise HTTPException(400, f'Vorschau nicht unterstützt für .{ext}')

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f'Lesefehler: {e}')

    truncated = len(text) > MAX_CHARS
    return {'text': text[:MAX_CHARS], 'truncated': truncated}


def _is_tar(path: str, fname: str) -> bool:
    f = fname.lower()
    return (f.endswith('.tar') or f.endswith('.tar.gz') or f.endswith('.tgz')
            or f.endswith('.tar.bz2') or f.endswith('.tar.xz'))

def _is_7z(fname: str) -> bool:
    return fname.lower().endswith('.7z')

try:
    import py7zr as _py7zr
    _HAS_7Z = True
except ImportError:
    _HAS_7Z = False


@router.get('/{att_id}/list')
async def list_archive(att_id: str, user_id: str = Depends(current_user_id)):
    with get_conn() as conn:
        att = conn.execute(
            'SELECT a.id, a.filename, a.path, a.card_id, c.board_id FROM attachments a JOIN cards c ON a.card_id=c.id WHERE a.id=?',
            (att_id,),
        ).fetchone()
        if not att:
            raise HTTPException(404, 'Attachment not found')
        require_board_access(att['board_id'], user_id, conn)

    if not os.path.exists(att['path']):
        raise HTTPException(404, 'Datei nicht gefunden')

    fname = att['filename'] or ''
    files = []

    if fname.lower().endswith('.zip'):
        try:
            with zipfile.ZipFile(att['path'], 'r') as zf:
                for info in zf.infolist():
                    files.append({
                        'name': info.filename,
                        'size': info.file_size,
                        'is_dir': info.is_dir(),
                    })
        except zipfile.BadZipFile:
            raise HTTPException(400, 'Ungültiges ZIP-Archiv')
    elif _is_tar(att['path'], fname):
        try:
            with tarfile.open(att['path'], 'r:*') as tf:
                for m in tf.getmembers():
                    files.append({
                        'name': m.name,
                        'size': m.size,
                        'is_dir': m.isdir(),
                    })
        except tarfile.TarError as e:
            raise HTTPException(400, f'Ungültiges Archiv: {e}')
    elif _is_7z(fname):
        if not _HAS_7Z:
            raise HTTPException(501, '7z-Unterstützung nicht verfügbar (py7zr fehlt)')
        try:
            with _py7zr.SevenZipFile(att['path'], 'r') as sz:
                for info in sz.list():
                    files.append({
                        'name': info.filename,
                        'size': info.uncompressed if info.uncompressed else 0,
                        'is_dir': info.is_directory,
                    })
        except Exception as e:
            raise HTTPException(400, f'Ungültiges 7z-Archiv: {e}')
    else:
        raise HTTPException(415, 'Format nicht unterstützt (ZIP, TAR, 7z)')

    return {'files': files, 'count': len(files)}


@router.post('/{att_id}/extract', status_code=201)
async def extract_zip(att_id: str, user_id: str = Depends(current_user_id)):
    import mimetypes
    with get_conn() as conn:
        att = conn.execute(
            '''SELECT a.id, a.filename, a.path, a.card_id, a.mime,
                      c.board_id, c.parent_card_id AS owner_parent, c.col_id, c.lane_id, c.card_type AS owner_type
               FROM attachments a JOIN cards c ON a.card_id=c.id WHERE a.id=?''',
            (att_id,),
        ).fetchone()
        if not att:
            raise HTTPException(404, 'Attachment not found')
        require_board_access(att['board_id'], user_id, conn)

    if not os.path.exists(att['path']):
        raise HTTPException(404, 'Datei nicht gefunden')

    fname_orig = att['filename'] or ''
    is_zip = fname_orig.lower().endswith('.zip')
    is_tar = _is_tar(att['path'], fname_orig)
    is_7z = _is_7z(fname_orig)
    if not is_zip and not is_tar and not is_7z:
        raise HTTPException(400, 'Nur ZIP, TAR und 7z können entpackt werden')
    if is_7z and not _HAS_7Z:
        raise HTTPException(501, '7z-Unterstützung nicht verfügbar (py7zr fehlt)')

    parent_card_id = att['owner_parent'] if (att['owner_parent'] and att['owner_type'] == 'file_card') else att['card_id']
    board_id = att['board_id']

    # Collect (filename, bytes) pairs
    entries: list[tuple[str, bytes]] = []
    if is_zip:
        try:
            with zipfile.ZipFile(att['path'], 'r') as zf:
                names = [n for n in zf.namelist() if not n.endswith('/') and not n.startswith('__MACOSX/')]
                if len(names) > 100:
                    raise HTTPException(400, 'Archiv enthält zu viele Dateien (max 100)')
                for name in names:
                    try:
                        data = zf.read(name)
                    except Exception:
                        continue
                    if len(data) <= MAX_SIZE:
                        entries.append((os.path.basename(name), data))
        except zipfile.BadZipFile:
            raise HTTPException(400, 'Ungültiges ZIP-Archiv')
    elif is_tar:
        try:
            with tarfile.open(att['path'], 'r:*') as tf:
                members = [m for m in tf.getmembers() if m.isfile() and not os.path.basename(m.name).startswith('.')]
                if len(members) > 100:
                    raise HTTPException(400, 'Archiv enthält zu viele Dateien (max 100)')
                for m in members:
                    try:
                        f = tf.extractfile(m)
                        if not f:
                            continue
                        data = f.read()
                    except Exception:
                        continue
                    if len(data) <= MAX_SIZE:
                        entries.append((os.path.basename(m.name), data))
        except tarfile.TarError as e:
            raise HTTPException(400, f'Ungültiges TAR-Archiv: {e}')
    else:  # 7z
        try:
            with _py7zr.SevenZipFile(att['path'], 'r') as sz:
                all_names = sz.getnames()
            file_names = [n for n in all_names if not n.endswith('/')]
            if len(file_names) > 100:
                raise HTTPException(400, 'Archiv enthält zu viele Dateien (max 100)')
            with tempfile.TemporaryDirectory() as tmpdir:
                with _py7zr.SevenZipFile(att['path'], 'r') as sz:
                    sz.extractall(path=tmpdir)
                for arc_name in file_names:
                    fpath = os.path.join(tmpdir, arc_name)
                    if not os.path.isfile(fpath):
                        continue
                    try:
                        with open(fpath, 'rb') as fh:
                            data = fh.read()
                    except Exception:
                        continue
                    if len(data) <= MAX_SIZE:
                        entries.append((os.path.basename(arc_name), data))
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(400, f'Ungültiges 7z-Archiv: {e}')

    created = []
    for fname, content in entries:
        if not fname:
            continue
        fc_id = uid()
        ts = now()
        with get_conn() as conn:
            parent = conn.execute('SELECT col_id, lane_id, board_id FROM cards WHERE id=?', (parent_card_id,)).fetchone()
            if not parent:
                continue
            pos = (conn.execute(
                'SELECT COALESCE(MAX(position),-1) AS m FROM cards WHERE parent_card_id=?', (parent_card_id,)
            ).fetchone()['m'] or -1) + 1
            conn.execute(
                '''INSERT INTO cards (id, board_id, col_id, lane_id, position, title, notes,
                   color, bg_color, cover_path, points, points_max,
                   card_type, parent_card_id, card_mode, time_spent, created_at, updated_at)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)''',
                (fc_id, board_id, parent['col_id'], parent['lane_id'], pos,
                 fname, '', None, None, None, 0, None, 'file_card', parent_card_id, 'org', 0, ts, ts),
            )
        card_dir = os.path.join(UPLOAD_PATH, fc_id)
        os.makedirs(card_dir, exist_ok=True)
        safe_name = _safe_filename(fname)
        att_id_new = uid()
        fpath = os.path.join(card_dir, f'{att_id_new}_{safe_name}')
        with open(fpath, 'wb') as f:
            f.write(content)
        import mimetypes as _mt
        mime, _ = _mt.guess_type(fname)
        with get_conn() as conn:
            conn.execute(
                'INSERT INTO attachments (id, card_id, filename, path, size, mime, created_at) VALUES (?,?,?,?,?,?,?)',
                (att_id_new, fc_id, fname, fpath, len(content), mime or 'application/octet-stream', ts),
            )
        created.append({'id': fc_id, 'title': fname})

    if created:
        broadcast(board_id, {'type': 'card_updated', 'userId': user_id, 'id': parent_card_id, 'boardId': board_id})

    return {'cards': created, 'count': len(created)}


@router.patch('/{att_id}/move')
def move_attachment(att_id: str, body: MoveAttachmentIn, user_id: str = Depends(current_user_id)):
    with get_conn() as conn:
        att = conn.execute(
            'SELECT a.*, c.board_id FROM attachments a JOIN cards c ON a.card_id=c.id WHERE a.id=?',
            (att_id,),
        ).fetchone()
        if not att:
            raise HTTPException(404, 'Attachment not found')
        old_board_id = att['board_id']
        old_card_id = att['card_id']
        require_board_access(old_board_id, user_id, conn)
        target = conn.execute('SELECT id, board_id FROM cards WHERE id=?', (body.card_id,)).fetchone()
        if not target:
            raise HTTPException(404, 'Zielkarte nicht gefunden')
        if target['board_id'] != old_board_id:
            raise HTTPException(400, 'Zielkarte ist in einem anderen Board')
        conn.execute('UPDATE attachments SET card_id=? WHERE id=?', (body.card_id, att_id))
        conn.execute('UPDATE cards SET updated_at=? WHERE id=?', (now(), old_card_id))
        conn.execute('UPDATE cards SET updated_at=? WHERE id=?', (now(), body.card_id))
    broadcast(old_board_id, {'type': 'card_updated', 'userId': user_id, 'id': old_card_id, 'boardId': old_board_id})
    broadcast(old_board_id, {'type': 'card_updated', 'userId': user_id, 'id': body.card_id, 'boardId': old_board_id})
    return {'ok': True}


class NoteToFilecardIn(BaseModel):
    notes: str = ''


@router.post('/note-to-filecard/{card_id}', status_code=201)
def note_to_filecard(card_id: str, body: NoteToFilecardIn, user_id: str = Depends(current_user_id)):
    import mimetypes as _mt
    from datetime import datetime
    notes = body.notes.strip()
    if not notes:
        raise HTTPException(400, 'Notiz ist leer')
    with get_conn() as conn:
        c = conn.execute('SELECT board_id, col_id, lane_id FROM cards WHERE id=?', (card_id,)).fetchone()
        if not c:
            raise HTTPException(404, 'Card not found')
        require_board_access(c['board_id'], user_id, conn)
        board_id, col_id, lane_id = c['board_id'], c['col_id'], c['lane_id']
        pos = (conn.execute(
            'SELECT COALESCE(MAX(position),-1) AS m FROM cards WHERE parent_card_id=?', (card_id,)
        ).fetchone()['m'] or -1) + 1

    ts = now()
    fname = 'notiz vom ' + datetime.now().strftime('%Y-%m-%d %H:%M') + '.md'
    fc_id = uid()
    att_id = uid()

    with get_conn() as conn:
        conn.execute(
            '''INSERT INTO cards (id, board_id, col_id, lane_id, position, title, notes,
               color, bg_color, cover_path, points, points_max,
               card_type, parent_card_id, card_mode, time_spent, created_at, updated_at)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)''',
            (fc_id, board_id, col_id, lane_id, pos,
             fname, '', None, None, None, 0, None, 'file_card', card_id, 'org', 0, ts, ts),
        )

    card_dir = os.path.join(UPLOAD_PATH, fc_id)
    os.makedirs(card_dir, exist_ok=True)
    safe_name = _safe_filename(fname)
    fpath = os.path.join(card_dir, f'{att_id}_{safe_name}')
    content = notes.encode('utf-8')
    with open(fpath, 'wb') as f:
        f.write(content)

    with get_conn() as conn:
        conn.execute(
            'INSERT INTO attachments (id, card_id, filename, path, size, mime, created_at) VALUES (?,?,?,?,?,?,?)',
            (att_id, fc_id, fname, fpath, len(content), 'text/markdown', ts),
        )

    broadcast(board_id, {'type': 'card_updated', 'userId': user_id, 'id': card_id, 'boardId': board_id})
    return {'id': fc_id}


# ── tul-DV Bridge (inter-service, X-Tul-Secret auth) ─────────────────────────

def _check_tul_secret(request: Request):
    # Eigenes Secret für die Hub<->mkan-Bridge, getrennt von tools_nuc's TUL_SECRET
    # (das dort zusätzlich als JWT-Signing-Key verwendet wird) — ein Leck hier soll
    # nicht auch Session-Tokens auf der tul-Seite kompromittieren. Fund Code-Review 2026-07-12.
    secret = os.environ.get('TUL_BRIDGE_SECRET', '')
    if not secret:
        raise HTTPException(503, 'TUL_BRIDGE_SECRET not configured')
    if not hmac.compare_digest(request.headers.get('X-Tul-Secret', ''), secret):
        raise HTTPException(401, 'Unauthorized')


@router.get('/dv/card/{card_id}')
def dv_card_files(card_id: str, request: Request):
    _check_tul_secret(request)
    with get_conn() as conn:
        card = conn.execute('SELECT id, title, dv_shared FROM cards WHERE id=?', (card_id,)).fetchone()
        if not card:
            raise HTTPException(404, 'Card not found')
        if not card['dv_shared']:
            raise HTTPException(403, 'DV sharing not enabled for this card')
        atts = conn.execute(
            'SELECT id, filename, size, mime FROM attachments WHERE card_id=? ORDER BY position, created_at',
            (card_id,),
        ).fetchall()
        fc_ids = [r['id'] for r in conn.execute(
            "SELECT id FROM cards WHERE parent_card_id=? AND card_type='file_card'", (card_id,),
        ).fetchall()]
        fc_atts = []
        for fcid in fc_ids:
            for r in conn.execute(
                'SELECT id, filename, size, mime FROM attachments WHERE card_id=?', (fcid,),
            ).fetchall():
                fc_atts.append(r)
    return {
        'card_id': card_id,
        'title': card['title'],
        'files': [
            {'id': a['id'], 'name': a['filename'], 'size': a['size'], 'mime': a['mime'] or ''}
            for a in list(atts) + fc_atts
        ],
    }


@router.get('/dv/file/{att_id}')
def dv_file(att_id: str, request: Request):
    _check_tul_secret(request)
    import mimetypes as _mt
    with get_conn() as conn:
        att = conn.execute(
            'SELECT a.*, c.dv_shared FROM attachments a JOIN cards c ON a.card_id=c.id WHERE a.id=?',
            (att_id,),
        ).fetchone()
        if not att:
            raise HTTPException(404, 'Attachment not found')
        if not att['dv_shared']:
            # Check if it's a file_card attachment whose parent is dv_shared
            att = conn.execute(
                '''SELECT a.*, pc.dv_shared
                   FROM attachments a
                   JOIN cards c ON a.card_id=c.id
                   JOIN cards pc ON c.parent_card_id=pc.id
                   WHERE a.id=? AND c.card_type='file_card' AND pc.dv_shared=1''',
                (att_id,),
            ).fetchone()
            if not att:
                raise HTTPException(403, 'DV sharing not enabled')
    if not os.path.exists(att['path']):
        raise HTTPException(404, 'File missing on disk')
    mime = att['mime'] or _mt.guess_type(att['filename'])[0] or 'application/octet-stream'
    return FileResponse(att['path'], filename=att['filename'], media_type=mime)


@router.get('/dv/pool')
def dv_pool(request: Request):
    _check_tul_secret(request)
    with get_conn() as conn:
        cards = conn.execute(
            "SELECT id, title FROM cards WHERE dv_shared=1 AND card_type='card' ORDER BY title"
        ).fetchall()
        result = []
        for card in cards:
            atts = conn.execute(
                'SELECT id, filename, size, mime FROM attachments WHERE card_id=? ORDER BY position, created_at',
                (card['id'],),
            ).fetchall()
            fc_ids = [r['id'] for r in conn.execute(
                "SELECT id FROM cards WHERE parent_card_id=? AND card_type='file_card'",
                (card['id'],),
            ).fetchall()]
            fc_atts = []
            for fcid in fc_ids:
                for r in conn.execute(
                    'SELECT id, filename, size, mime FROM attachments WHERE card_id=?', (fcid,),
                ).fetchall():
                    fc_atts.append(r)
            all_files = list(atts) + fc_atts
            if all_files:
                result.append({
                    'card_id': card['id'],
                    'card_title': card['title'],
                    'files': [
                        {'id': a['id'], 'name': a['filename'], 'size': a['size'], 'mime': a['mime'] or ''}
                        for a in all_files
                    ],
                })
    return {'pool': result}


@router.get('/dv/cards-for-tool/{tul}')
def dv_cards_for_tool(tul: str, request: Request):
    _check_tul_secret(request)
    import json as _json
    with get_conn() as conn:
        cards = conn.execute(
            "SELECT id, title, card_settings FROM cards WHERE dv_shared=1 AND card_type='card'",
        ).fetchall()
        result = []
        for card in cards:
            try:
                cs = _json.loads(card['card_settings'] or '{}')
            except Exception:
                cs = {}
            if cs.get('targetTul') != tul:
                continue
            atts = conn.execute(
                'SELECT id, filename, size, mime FROM attachments WHERE card_id=? ORDER BY position, created_at',
                (card['id'],),
            ).fetchall()
            fc_ids = [r['id'] for r in conn.execute(
                "SELECT id FROM cards WHERE parent_card_id=? AND card_type='file_card'", (card['id'],),
            ).fetchall()]
            fc_atts = []
            for fcid in fc_ids:
                for r in conn.execute(
                    'SELECT id, filename, size, mime FROM attachments WHERE card_id=?', (fcid,),
                ).fetchall():
                    fc_atts.append(r)
            result.append({
                'card_id': card['id'],
                'title': card['title'],
                'files': [
                    {'id': a['id'], 'name': a['filename'], 'size': a['size'], 'mime': a['mime'] or ''}
                    for a in list(atts) + fc_atts
                ],
            })
    return result


@router.get('/dv/preview-text/{att_id}')
def dv_preview_text(att_id: str, request: Request):
    _check_tul_secret(request)
    with get_conn() as conn:
        att = conn.execute(
            '''SELECT a.id, a.filename, a.path, c.dv_shared
               FROM attachments a JOIN cards c ON a.card_id=c.id WHERE a.id=?''',
            (att_id,),
        ).fetchone()
        if not att:
            raise HTTPException(404, 'Attachment not found')
        if not att['dv_shared']:
            att = conn.execute(
                '''SELECT a.id, a.filename, a.path, pc.dv_shared
                   FROM attachments a JOIN cards c ON a.card_id=c.id
                   JOIN cards pc ON c.parent_card_id=pc.id
                   WHERE a.id=? AND c.card_type='file_card' AND pc.dv_shared=1''',
                (att_id,),
            ).fetchone()
            if not att:
                raise HTTPException(403, 'DV sharing not enabled')

    path = att['path']
    name = att['filename'] or ''
    ext = name.rsplit('.', 1)[-1].lower() if '.' in name else ''
    MAX_CHARS = 3000

    try:
        if ext == 'docx':
            from docx import Document
            doc = Document(path)
            text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        elif ext == 'odt':
            import xml.etree.ElementTree as ET
            NS_T = 'urn:oasis:names:tc:opendocument:xmlns:text:1.0'
            with zipfile.ZipFile(path) as z:
                with z.open('content.xml') as f:
                    tree = ET.parse(f)
            text = '\n'.join(
                t for t in (''.join(p.itertext()) for p in tree.iter(f'{{{NS_T}}}p')) if t.strip()
            )
        elif ext == 'xlsx':
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            ws = wb.active
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else '' for c in row]
                if any(c.strip() for c in cells):
                    rows.append('\t'.join(cells))
                if sum(len(r) for r in rows) > MAX_CHARS:
                    break
            wb.close()
            text = '\n'.join(rows)
        elif ext == 'ods':
            import xml.etree.ElementTree as ET
            NS_TAB = 'urn:oasis:names:tc:opendocument:xmlns:table:1.0'
            with zipfile.ZipFile(path) as z:
                with z.open('content.xml') as f:
                    tree = ET.parse(f)
            rows = []
            for row in tree.iter(f'{{{NS_TAB}}}table-row'):
                cells = [''.join(cell.itertext()).strip()
                         for cell in row.iter(f'{{{NS_TAB}}}table-cell')]
                line = '\t'.join(cells).rstrip('\t')
                if line.strip():
                    rows.append(line)
                if sum(len(r) for r in rows) > MAX_CHARS:
                    break
            text = '\n'.join(rows)
        elif ext == 'pptx':
            import xml.etree.ElementTree as ET
            NS_A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
            parts = []
            with zipfile.ZipFile(path) as z:
                slides = sorted(n for n in z.namelist()
                                if n.startswith('ppt/slides/slide') and n.endswith('.xml'))
                for sname in slides:
                    with z.open(sname) as f:
                        tree = ET.parse(f)
                    texts = [t.text for t in tree.iter(f'{{{NS_A}}}t')
                             if t.text and t.text.strip()]
                    if texts:
                        parts.append(' '.join(texts))
                    if sum(len(p) for p in parts) > MAX_CHARS:
                        break
            text = '\n'.join(parts)
        elif ext == 'odp':
            import xml.etree.ElementTree as ET
            NS_T = 'urn:oasis:names:tc:opendocument:xmlns:text:1.0'
            with zipfile.ZipFile(path) as z:
                with z.open('content.xml') as f:
                    tree = ET.parse(f)
            text = '\n'.join(
                t for t in (''.join(p.itertext()) for p in tree.iter(f'{{{NS_T}}}p')) if t.strip()
            )
        elif ext == 'xls':
            import xlrd
            wb = xlrd.open_workbook(path)
            ws = wb.sheet_by_index(0)
            rows = []
            for i in range(ws.nrows):
                line = '\t'.join(str(ws.cell_value(i, j)) for j in range(ws.ncols) if ws.cell_value(i, j))
                if line.strip():
                    rows.append(line)
                if sum(len(r) for r in rows) > MAX_CHARS:
                    break
            text = '\n'.join(rows)
        elif ext == 'doc':
            import subprocess
            result = subprocess.run(['antiword', path], capture_output=True, timeout=15)
            if result.returncode != 0:
                raise HTTPException(400, 'antiword nicht verfügbar oder Lesefehler')
            text = result.stdout.decode('utf-8', errors='replace')
        else:
            raise HTTPException(400, f'Vorschau nicht unterstützt für .{ext}')
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f'Lesefehler: {e}')

    truncated = len(text) > MAX_CHARS
    return {'text': text[:MAX_CHARS], 'truncated': truncated}


@router.post('/dv/unlink-card/{card_id}')
def dv_unlink_card(card_id: str, request: Request):
    _check_tul_secret(request)
    import json as _json
    with get_conn() as conn:
        card = conn.execute('SELECT id, card_settings, board_id FROM cards WHERE id=?', (card_id,)).fetchone()
        if not card:
            raise HTTPException(404, 'Card not found')
        try:
            cs = _json.loads(card['card_settings'] or '{}')
        except Exception:
            cs = {}
        cs.pop('targetTul', None)
        conn.execute('UPDATE cards SET card_settings=? WHERE id=?', (_json.dumps(cs), card_id))
        board_id = card['board_id']
    broadcast(board_id, {'type': 'card_updated', 'id': card_id, 'boardId': board_id})
    return {'ok': True}


@router.post('/dv/upload-to-card/{card_id}', status_code=201)
async def dv_upload_to_card(card_id: str, request: Request, file: UploadFile = File(...)):
    _check_tul_secret(request)
    with get_conn() as conn:
        card = conn.execute(
            'SELECT id, dv_shared, board_id, col_id, lane_id FROM cards WHERE id=?', (card_id,)
        ).fetchone()
        if not card:
            raise HTTPException(404, 'Card not found')
        if not card['dv_shared']:
            raise HTTPException(403, 'DV sharing not enabled for this card')
        board_id = card['board_id']
        col_id = card['col_id']
        lane_id = card['lane_id']
        pos = (conn.execute(
            "SELECT COALESCE(MAX(position),-1) AS m FROM cards WHERE parent_card_id=? AND card_type='file_card'",
            (card_id,)
        ).fetchone()['m'] or -1) + 1

    fname = file.filename or 'upload'
    safe_name = _safe_filename(fname)
    fc_id = uid()
    att_id = uid()
    ts = now()

    card_dir = os.path.join(UPLOAD_PATH, fc_id)
    os.makedirs(card_dir, exist_ok=True)
    path = os.path.join(card_dir, f'{att_id}_{safe_name}')
    size = await _stream_to_file(file, path)
    mime = file.content_type or ''

    with get_conn() as conn:
        conn.execute(
            '''INSERT INTO cards (id, board_id, col_id, lane_id, position, title, notes,
               color, bg_color, cover_path, points, points_max,
               card_type, parent_card_id, card_mode, time_spent, created_at, updated_at)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)''',
            (fc_id, board_id, col_id, lane_id, pos,
             fname, '', None, None, None, 0, None, 'file_card', card_id, 'org', 0, ts, ts),
        )
        conn.execute(
            'INSERT INTO attachments (id, card_id, filename, path, size, mime, created_at) VALUES (?,?,?,?,?,?,?)',
            (att_id, fc_id, fname, path, size, mime, ts),
        )
    broadcast(board_id, {'type': 'card_updated', 'id': card_id, 'boardId': board_id})
    return {'id': fc_id, 'name': fname, 'size': size}
